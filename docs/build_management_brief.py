from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = r"D:\Log Files\docs\Express_Intelligence_500GB_Management_Brief.docx"
NAVY = "0B2545"
BLUE = "2E74B5"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
MID_GRAY = "68717D"
RED = "B91C2B"
WHITE = "FFFFFF"
INK = "20252B"


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa):
    total = sum(widths_dxa)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(widths_dxa[idx] / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell_margins(cell)


def set_repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def font_run(run, size=11, bold=False, color=INK, name="Calibri", italic=False):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def add_bullet(doc, text, bold_lead=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    if bold_lead and text.startswith(bold_lead):
        font_run(p.add_run(bold_lead), bold=True)
        font_run(p.add_run(text[len(bold_lead):]))
    else:
        font_run(p.add_run(text))
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    font_run(p.add_run(text))
    return p


def add_callout(doc, label, text, fill=LIGHT_BLUE):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    shade(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    font_run(p.add_run(label + "  "), size=10, bold=True, color=NAVY)
    font_run(p.add_run(text), size=10.5, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        shade(cell, NAVY)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        font_run(p.add_run(header), size=9.5, bold=True, color=WHITE)
    set_repeat_header(table.rows[0])
    for row_idx, values in enumerate(rows):
        cells = table.add_row().cells
        if row_idx % 2:
            for cell in cells:
                shade(cell, LIGHT_GRAY)
        for i, value in enumerate(values):
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.08
            font_run(p.add_run(value), size=9.5)
    set_table_geometry(table, widths)
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(2)
    return table


def page_break(doc):
    doc.add_page_break()


doc = Document()
sec = doc.sections[0]
sec.page_width = Inches(8.5)
sec.page_height = Inches(11)
sec.top_margin = Inches(0.78)
sec.bottom_margin = Inches(0.72)
sec.left_margin = Inches(1)
sec.right_margin = Inches(1)
sec.header_distance = Inches(0.492)
sec.footer_distance = Inches(0.492)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
normal.font.size = Pt(11)
normal.font.color.rgb = RGBColor.from_string(INK)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.25
for name, size, color, before, after in (
    ("Heading 1", 16, BLUE, 18, 10),
    ("Heading 2", 13, BLUE, 14, 7),
    ("Heading 3", 12, NAVY, 10, 5),
):
    style = styles[name]
    style.font.name = "Calibri"
    style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = RGBColor.from_string(color)
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.keep_with_next = True
for list_name in ("List Bullet", "List Number"):
    style = styles[list_name]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.paragraph_format.left_indent = Inches(0.375)
    style.paragraph_format.first_line_indent = Inches(-0.188)
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.line_spacing = 1.25

# Running header and footer
hp = sec.header.paragraphs[0]
hp.text = "EXPRESS INTELLIGENCE OS  |  MANAGEMENT BRIEF"
hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
font_run(hp.runs[0], size=8.5, bold=True, color=MID_GRAY)
fp = sec.footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
font_run(fp.add_run("Internal planning  |  25 August 2026  |  Page "), size=8.5, color=MID_GRAY)
fld = OxmlElement("w:fldSimple")
fld.set(qn("w:instr"), "PAGE")
fp._p.append(fld)

# Page 1: memo masthead
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(18)
p.paragraph_format.space_after = Pt(2)
font_run(p.add_run("MANAGEMENT DECISION BRIEF"), size=10, bold=True, color=RED)
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(5)
font_run(p.add_run("Scaling Express Intelligence OS"), size=27, bold=True, color=NAVY)
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(18)
font_run(p.add_run("A practical operating stack for daily CDN and origin log analysis up to 500 GB"), size=14, color=MID_GRAY)
add_table(doc, ["Decision", "Recommended direction"], [
    ("Hosting", "Vercel for the team interface; private India-region compute for uploads and analysis."),
    ("Data", "Keep raw logs and derived evidence on company-controlled encrypted storage."),
    ("Scale", "Start with a 2 TB NVMe active tier and expand storage separately from compute."),
    ("Release gate", "Do not promise 500 GB processing time until a representative load test passes."),
], [2160, 7200])
add_callout(doc, "BOTTOM LINE", "Vercel alone cannot process or retain 500 GB log files. The browser interface can remain on Vercel, but the API, workers, database, Redis and storage must run on durable infrastructure.")
doc.add_heading("What management is approving", level=1)
for text in [
    "A private production environment that team members reach through one HTTPS URL.",
    "Enough fast local storage for active analysis, plus separate backup or archive capacity.",
    "A measured rollout: 25 GB, 100 GB, then 500 GB acceptance tests before daily production use.",
    "Named operational ownership for access, backups, monitoring, upgrades and incident response.",
]: add_bullet(doc, text)
doc.add_heading("What this plan does not claim", level=2)
for text in [
    "It does not guarantee that every 500 GB archive will finish in a fixed number of hours.",
    "It does not make a laptop or free hosting tier suitable for continuous team production.",
    "It does not replace source-file retention policy, legal review or company security approval.",
]: add_bullet(doc, text)

page_break(doc)
doc.add_heading("1. The stack in plain language", level=1)
add_table(doc, ["Layer", "Purpose", "Recommended component"], [
    ("Team interface", "Where users upload, monitor and review", "Next.js on Vercel"),
    ("Secure entry", "HTTPS, request limits and controlled access", "Caddy or managed load balancer"),
    ("Application API", "Accepts jobs and serves results", "FastAPI in Docker"),
    ("Background workers", "Stream, decompress and aggregate logs", "Python workers in Docker"),
    ("Job coordination", "Durable queue and progress state", "Redis with persistence"),
    ("Product database", "Runs, metrics, evidence and audit history", "PostgreSQL"),
    ("Active file storage", "Raw files and working data", "Encrypted NVMe volume"),
    ("Backup/archive", "Recovery and retention", "Separate encrypted disk or India-region object storage"),
    ("Operations", "Health, capacity, errors and alerts", "Metrics, logs, alerts and daily backup checks"),
], [1500, 3220, 4640])
add_callout(doc, "DATA PATH", "User -> Vercel interface -> HTTPS API -> persistent storage -> worker -> PostgreSQL evidence -> results dashboard. The worker reads streams; it must not load the whole file into memory.")
doc.add_heading("Why these parts are separate", level=2)
for text in [
    "Compute can be increased without moving the interface.",
    "Storage can grow without redesigning the entire application.",
    "A failed browser session does not cancel a background job.",
    "The database stores durable evidence, while Redis coordinates temporary work.",
]: add_bullet(doc, text)

page_break(doc)
doc.add_heading("2. Capacity required for heavy daily files", level=1)
add_table(doc, ["Operating tier", "Compute", "Active storage", "Best use"], [
    ("Pilot: up to 25 GB", "8-16 CPU cores, 32 GB RAM", "500 GB NVMe", "Parser validation and team workflow"),
    ("Standard: up to 100 GB", "16-32 cores, 64 GB RAM", "1 TB NVMe", "Regular production with one large job"),
    ("Heavy: up to 500 GB", "32-64 cores, 128 GB RAM", "2 TB+ NVMe", "Controlled 500 GB runs after load testing"),
], [1720, 2420, 2280, 2940])
add_callout(doc, "STORAGE RULE", "Reserve space for the source file, temporary working data, database growth and recovery. A 500 GB input does not fit safely on a disk with only 500 GB free. Begin with at least 2 TB active NVMe and monitor the actual expansion ratio.", fill="FFF3D6")
doc.add_heading("Minimum guardrails", level=2)
for text in [
    "One active upload per user; cap total active uploads for the server.",
    "Reject a job before upload when free space is below the configured reserve.",
    "Chunked, resumable uploads with checksums and clear retry status.",
    "Stream ZIP members and reject unsafe paths, excessive member counts and extreme expansion ratios.",
    "Bound worker concurrency so CPU, memory, disk and PostgreSQL remain responsive.",
    "Keep raw files, scratch data, PostgreSQL and backups on deliberate persistent volumes.",
]: add_bullet(doc, text)
doc.add_heading("Retention questions management must decide", level=2)
for text in [
    "How many days are raw CDN and origin files retained?",
    "How many completed result sets remain online?",
    "Who can delete a source file or analysis result?",
    "Is archive storage required in India, and what recovery time is acceptable?",
]: add_bullet(doc, text)

page_break(doc)
doc.add_heading("3. Deployment choices and cost categories", level=1)
add_table(doc, ["Option", "Advantages", "Trade-offs"], [
    ("Company server / data centre", "Maximum control; predictable local storage; no public raw-data transfer", "Company must operate hardware, power, backups, patching and remote access"),
    ("India-region cloud VM", "Fastest shared rollout; easier resizing; managed network controls", "Monthly compute, disk, backup and outbound-transfer charges"),
    ("Hybrid: Vercel + private backend", "Simple team URL; independent backend scaling; recommended architecture", "Requires secure HTTPS connection and operational ownership"),
], [1940, 3620, 3800])
doc.add_heading("Budget categories", level=2)
for text in [
    "Compute: API, workers and database CPU/RAM.",
    "Storage: active NVMe, database volume, snapshots and archive retention.",
    "Network: large-file ingress, result access and any outbound transfer.",
    "Reliability: monitoring, backup verification and replacement capacity.",
    "Security: access controls, vulnerability patching, secrets and audit review.",
    "People: an infrastructure owner and an application owner.",
]: add_bullet(doc, text, bold_lead=text.split(":")[0] + ":")
add_callout(doc, "FREE SOFTWARE IS NOT FREE INFRASTRUCTURE", "Docker, PostgreSQL, Redis, FastAPI and Caddy have open-source editions. Reliable 500 GB daily processing still needs paid or company-owned compute, storage, backup and operations.")
doc.add_heading("Procurement request", level=2)
for text in [
    "Ask vendors for India-region pricing using 2 TB active NVMe, 64-128 GB RAM and separate backups.",
    "Request monthly estimates at pilot, 100 GB and 500 GB tiers.",
    "Do not select solely on headline VM cost; include disk performance, snapshots and data transfer.",
]: add_bullet(doc, text)

page_break(doc)
doc.add_heading("4. Security and team access", level=1)
doc.add_heading("Required before external team use", level=2)
for text in [
    "Use HTTPS only; never expose PostgreSQL or Redis to the public internet.",
    "Give every user an individual identity. Remove access immediately when a person leaves the team.",
    "Restrict uploads by size, type, rate and concurrent-job capacity.",
    "Store passwords, database URLs and service credentials outside Git.",
    "Encrypt disks and backups; limit filesystem access to the application services.",
    "Record uploads, job starts, exports, failures and administrative changes in an audit trail.",
    "Patch base images and dependencies on a defined schedule; scan releases before production.",
    "Test backup restore, not only backup creation.",
]: add_bullet(doc, text)
add_callout(doc, "PRIVACY BOUNDARY", "Vercel should host interface assets and browser code. Raw logs should travel only to the approved private backend and remain in company-controlled storage.")
doc.add_heading("Suggested access levels", level=2)
add_table(doc, ["Access", "Can do"], [
    ("Analyst", "Upload approved files, start analysis, view and export permitted results"),
    ("Operator", "Manage queues, retry failed jobs, check storage and perform recovery"),
    ("Administrator", "Manage users, configuration, retention and audit review"),
], [1900, 7460])
doc.add_heading("Incident basics", level=2)
for text in [
    "Pause new uploads when disk reserve, database health or error thresholds fail.",
    "Preserve logs and audit events; do not overwrite evidence during investigation.",
    "Rotate affected credentials and document the recovery decision.",
]: add_bullet(doc, text)

page_break(doc)
doc.add_heading("5. Performance expectations", level=1)
add_callout(doc, "NO FIXED PROMISE BEFORE TESTING", "Processing time depends on compressed size, expanded size, record count, parser acceptance, URL cardinality, disk speed and worker concurrency. The product should show an estimate range and confidence level after sampling the uploaded file.", fill="FFF3D6")
doc.add_heading("What the user should see before starting", level=2)
for text in [
    "Selected analysis amount: sample rows, file percentage, exact GB limit or full file.",
    "Estimated records and expanded bytes based on preflight sampling.",
    "Estimated completion range, for example 2-4 hours, not a false exact minute.",
    "Estimate confidence: low, medium or high, with the assumptions shown.",
    "Required free storage and current free storage.",
    "A warning when the estimate exceeds the normal operating window.",
]: add_bullet(doc, text)
doc.add_heading("How estimates become reliable", level=2)
for text in [
    "Measure sample throughput for reading, decompression, parsing and aggregation.",
    "Use recent completed runs with the same source type and similar hardware.",
    "Update the remaining-time range while the job runs.",
    "Keep a visible difference between upload progress and analysis progress.",
]: add_number(doc, text)
doc.add_heading("Acceptance targets", level=2)
add_table(doc, ["Test", "Pass condition"], [
    ("Recovery", "Restart services during a job; it resumes or safely retries without duplicate results"),
    ("Capacity", "Disk reserve is enforced before and during upload"),
    ("Correctness", "Processed + rejected counts reconcile with source evidence"),
    ("Performance", "25 GB, 100 GB and 500 GB benchmarks meet an agreed operating window"),
    ("Usability", "Users understand selected scope, estimate, progress, failure and next action"),
], [1900, 7460])

page_break(doc)
doc.add_heading("6. Operating checklist", level=1)
doc.add_heading("Before launch", level=2)
for text in [
    "Approve hosting location, retention period, owners and recovery objective.",
    "Provision HTTPS, private networking, encrypted volumes and separate backups.",
    "Run 25 GB and 100 GB load tests; tune worker and database limits.",
    "Run the 500 GB test using representative CDN and origin archives.",
    "Verify restore, restart, duplicate prevention and low-disk behaviour.",
    "Publish a user guide and an escalation contact.",
]: add_bullet(doc, text)
doc.add_heading("Every day", level=2)
for text in [
    "Check API, worker, PostgreSQL and Redis health.",
    "Check active-storage free space and backup completion.",
    "Review stuck, failed or unusually slow runs.",
    "Confirm completed results are visible after browser or system restart.",
]: add_bullet(doc, text)
doc.add_heading("Every week", level=2)
for text in [
    "Review throughput trends, rejected rows and estimate accuracy.",
    "Apply approved security updates and review audit events.",
    "Delete or archive data according to retention policy.",
]: add_bullet(doc, text)
doc.add_heading("Management approval checklist", level=2)
add_table(doc, ["Approval", "Owner / status"], [
    ("India-region hosting and budget", "____________________________"),
    ("Raw-log retention and deletion policy", "____________________________"),
    ("Security and access model", "____________________________"),
    ("Backup and recovery target", "____________________________"),
    ("500 GB performance acceptance window", "____________________________"),
    ("Production owner and support contact", "____________________________"),
], [5200, 4160])
add_callout(doc, "RECOMMENDED NEXT DECISION", "Approve the hybrid architecture and a staged capacity test. Treat 500 GB as a verified production tier only after the representative benchmark, restart and recovery checks pass.")

doc.core_properties.title = "Express Intelligence OS - 500 GB Management Brief"
doc.core_properties.subject = "Infrastructure, security and operations for heavy daily log analysis"
doc.core_properties.author = "Express Intelligence OS"
doc.core_properties.keywords = "log analysis, 500 GB, infrastructure, security, operations"
doc.save(OUT)
print(OUT)
